"""
Convert CLAUDE_API_BUDGET_REQUEST.md → CLAUDE_API_BUDGET_REQUEST.docx
using python-docx. Handles: headings, paragraphs, tables, lists, bold,
inline code, and horizontal rules.

Usage:
    python scripts/md_to_docx.py
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
# Default to the budget doc, or accept a custom .md path as first argument.
_default_src = ROOT / "CLAUDE_API_BUDGET_REQUEST.md"
SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else _default_src
DST = SRC.with_suffix(".docx")


# ---------- inline parser ----------

_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"          # bold
    r"|(`[^`]+`)"                # inline code
    r"|(\[[^\]]+\]\([^)]+\))"    # link
)


def add_inline_runs(paragraph, text):
    """Walk text, splitting on bold / code / link spans."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            label_end = token.index("]")
            label = token[1:label_end]
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x18, 0x77, 0xF2)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------- table helpers ----------

def parse_table(lines, idx):
    """Parse a markdown table starting at lines[idx]. Returns (rows, next_idx)."""
    header = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
    # separator line at idx+1
    rows = []
    i = idx + 2
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return header, rows, i


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for j, label in enumerate(header):
        p = hdr_cells[j].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        _shade_cell(hdr_cells[j], "1877F2")
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row in enumerate(rows, start=1):
        row_cells = table.rows[ri].cells
        for cj, cell_text in enumerate(row):
            if cj >= len(row_cells):
                continue
            p = row_cells[cj].paragraphs[0]
            add_inline_runs(p, cell_text)
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# ---------- main converter ----------

def convert(md_path, out_path):
    doc = Document()

    # Default style tuning
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    in_list = False

    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "BFBFBF")
            pbdr.append(bottom)
            p_pr.append(pbdr)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1C, 0x2B, 0x33)
            i += 1
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x18, 0x77, 0xF2)
            i += 1
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1C, 0x2B, 0x33)
            i += 1
            continue

        # Table
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            header, rows, next_i = parse_table(lines, i)
            add_table(doc, header, rows)
            i = next_i
            continue

        # Bulleted list item
        if re.match(r"^\s*-\s+", line):
            content = re.sub(r"^\s*-\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            # Handle checkbox lists "[ ]" or "[x]"
            checkbox_match = re.match(r"^\[([ xX])\]\s*(.+)$", content)
            if checkbox_match:
                state = checkbox_match.group(1).strip().lower() == "x"
                content = ("☑ " if state else "☐ ") + checkbox_match.group(2)
            add_inline_runs(p, content)
            i += 1
            continue

        # Numbered list item
        if re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, content)
            i += 1
            continue

        # Bold-only line (used as section emphasis)
        # Blank line
        if not line.strip():
            i += 1
            continue

        # Fenced code block
        if line.lstrip().startswith("```"):
            lang = line.strip().lstrip("`")
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Blockquote
        if line.lstrip().startswith("> "):
            content = line.lstrip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x65, 0x67, 0x6B)
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    if not SRC.exists():
        print(f"ERROR: Source markdown not found at {SRC}")
        sys.exit(1)
    convert(SRC, DST)

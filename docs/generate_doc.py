from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "D1D5DB"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_style(para, size=11, bold=False, color="1F2937", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc, text, size=11, bold=False, color="1F2937",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.color.rgb = RGBColor.from_string(color)
    p.alignment      = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_rule(doc, color="E5E7EB"):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

# ══════════════════════════════════════════════════════════════
#  HEADER BLOCK
# ══════════════════════════════════════════════════════════════
# Blue accent bar — simulate with a shaded 1-col table
header_tbl = doc.add_table(rows=1, cols=1)
header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hcell = header_tbl.rows[0].cells[0]
set_cell_bg(hcell, "1877F2")
hcell.width = Inches(6.5)
hp = hcell.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_before = Pt(14)
hp.paragraph_format.space_after  = Pt(14)
hr = hp.add_run("  AdAuto Dashboard — Unpopulated Report Fields")
hr.font.size      = Pt(18)
hr.font.bold      = True
hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()  # spacer

# Sub-title
add_para(doc, "Technical Explanation & Resolution Roadmap",
         size=13, bold=False, color="6B7280", space_after=4)

# Meta info table
meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["Document Version", "Prepared by", "Date", "Status"]
values = ["1.0", "AdAuto Engineering", datetime.date.today().strftime("%B %Y"), "Open — Pending Resolution"]
for i, (lbl, val) in enumerate(zip(labels, values)):
    lc = meta.rows[i].cells[0]
    vc = meta.rows[i].cells[1]
    lc.width = Inches(1.8)
    vc.width = Inches(4.0)
    lp = lc.paragraphs[0]
    lr = lp.add_run(lbl)
    lr.font.size = Pt(9); lr.font.bold = True; lr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    vp = vc.paragraphs[0]
    vr = vp.add_run(val)
    vr.font.size = Pt(9); vr.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

doc.add_paragraph()
add_rule(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  OVERVIEW
# ══════════════════════════════════════════════════════════════
add_para(doc, "Overview", size=14, bold=True, color="111827", space_before=4, space_after=6)
add_para(doc,
    "This document explains why certain fields in the Weekly Performance Report are currently unpopulated. "
    "Each field is categorised by its root cause, the technical constraint involved, and the steps required to resolve it.",
    size=10.5, color="374151", space_after=10)

# Category legend table
leg = doc.add_table(rows=3, cols=2)
leg.alignment = WD_TABLE_ALIGNMENT.LEFT
leg.style = "Table Grid"
# Header row
set_cell_bg(leg.rows[0].cells[0], "F3F4F6")
set_cell_bg(leg.rows[0].cells[1], "F3F4F6")
for ci, txt in enumerate(["Category", "Meaning"]):
    cp = leg.rows[0].cells[ci].paragraphs[0]
    cr = cp.add_run(txt)
    cr.font.bold = True; cr.font.size = Pt(9.5); cr.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

rows_data = [
    ("⚠️  Partial",     "Data is available but requires additional conventions or approximations"),
    ("❌  Unavailable", "Data is not currently being collected — requires a system change"),
]
for ri, (cat, meaning) in enumerate(rows_data, start=1):
    for ci, txt in enumerate([cat, meaning]):
        cp = leg.rows[ri].cells[ci].paragraphs[0]
        cr = cp.add_run(txt)
        cr.font.size = Pt(9.5); cr.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

for row in leg.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
        cell._tc.get_or_add_tcPr()

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  SECTION HELPER
# ══════════════════════════════════════════════════════════════
def section_heading(doc, label, color_hex="FEF3C7", text_hex="92400E", prefix="⚠️  Partial Fields"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"  {label}")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(text_hex)
    doc.add_paragraph()

def field_block(doc, number, title, why, requirement, action_owner, effort, status, action_required):
    # Field title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{number}.  ")
    r1.font.size = Pt(12); r1.font.bold = True; r1.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    r2 = p.add_run(title)
    r2.font.size = Pt(12); r2.font.bold = True; r2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    def labeled_block(label, body, label_color="6B7280", body_color="374151"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.8)
        rl = p.add_run(label + "  ")
        rl.font.size = Pt(9.5); rl.font.bold = True
        rl.font.color.rgb = RGBColor.from_string(label_color)
        rb = p.add_run(body)
        rb.font.size = Pt(9.5)
        rb.font.color.rgb = RGBColor.from_string(body_color)

    labeled_block("Why it is not populated:", why)
    labeled_block("What is required:", requirement)
    labeled_block("Action owner:", action_owner)
    labeled_block("Estimated effort:", effort)
    labeled_block("Current status:", status)
    labeled_block("Action required:", action_required)

    add_rule(doc, "E5E7EB")
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  PARTIAL FIELDS
# ══════════════════════════════════════════════════════════════
section_heading(doc, "⚠️  Partial Fields", color_hex="FFFBEB", text_hex="92400E")
add_para(doc,
    "These fields can be populated, but only under specific conditions or with an acknowledged approximation.",
    size=10, color="78716C", space_after=10)

field_block(
    doc, 1,
    "Retargeting CPP & Broad Audience CPP",
    "The Meta Ads API does not expose a native audience-type label per ad set. There is no field that explicitly "
    "marks an ad set as retargeting versus broad prospecting. This distinction exists only in the naming convention "
    "applied by the media buyer when the ad set is created.",
    "Ad set names stored in the database must follow a consistent naming convention so the system can classify them. "
    "For example — Retargeting: [RET] Website Visitors 30D  |  Broad: [BROAD] India 25-45 Interest. "
    "Once conventions are enforced, the system pattern-matches on adset_name to segment CPP accordingly.",
    "Media Buyer",
    "Low — no engineering work needed once naming is applied",
    "Not populated. Will be enabled once naming conventions are consistently applied across all active campaigns.",
    "Rename all active ad sets to include a clear prefix: RET, BROAD, LAL (Lookalike), or COLD."
)

field_block(
    doc, 2,
    "CBO Budget Distribution — Top Ad Set %",
    "Campaign Budget Optimisation allocates budget dynamically. The Meta API does not return a planned budget share "
    "per ad set — because CBO is designed to be dynamic and does not expose a fixed allocation percentage.",
    "Derive spend distribution from actual spend per ad set using: Ad Set Spend Share (%) = Ad Set Spend ÷ Total "
    "Campaign Spend × 100. This is an approximation based on actual spend, not planned budget allocation.",
    "Engineering",
    "Low — 1–2 hours",
    "Derivable as an approximation. Not shown by default to avoid misrepresenting it as a planned budget figure.",
    "Add a spend-distribution view with a clear 'Actual Spend Share' label and disclaimer in the dashboard."
)

# ══════════════════════════════════════════════════════════════
#  UNAVAILABLE FIELDS
# ══════════════════════════════════════════════════════════════
section_heading(doc, "❌  Unavailable Fields", color_hex="FEF2F2", text_hex="991B1B")
add_para(doc,
    "These fields cannot be populated until a change is made to how data is collected from the Meta Ads API.",
    size=10, color="78716C", space_after=10)

field_block(
    doc, 3,
    "Ad Sets Paused / Killed This Week",
    "The current data model only stores performance metrics (spend, impressions, clicks, etc.) on a per-day basis. "
    "It does not store the status of an ad set (Active, Paused, Archived) at any point in time. "
    "To detect whether an ad set was paused or killed during a given week, the system would need to compare "
    "status between consecutive days — data that has never been collected.",
    "Create a new database table adset_status_log with columns: adset_id, adset_name, campaign_id, status, "
    "recorded_date. During each Meta sync, request ad set statuses from the API and insert a daily snapshot. "
    "In the report, query for ad sets where status changed from ACTIVE to PAUSED or DELETED within the week.",
    "Engineering",
    "Medium — 3–4 hours including schema migration, sync update, and frontend display",
    "No historical status data exists. Field shows — until tracking is implemented.",
    "Add adset status snapshotting to the sync pipeline as a future sprint item."
)

# ══════════════════════════════════════════════════════════════
#  SUMMARY TABLE
# ══════════════════════════════════════════════════════════════
add_para(doc, "Summary", size=14, bold=True, color="111827", space_before=8, space_after=8)

headers = ["Field", "Category", "Root Cause", "Action Owner", "Effort"]
rows    = [
    ["Retargeting CPP",          "⚠️ Partial",       "No naming convention on ad sets",                    "Media Buyer", "Low"],
    ["Broad Audience CPP",       "⚠️ Partial",       "No naming convention on ad sets",                    "Media Buyer", "Low"],
    ["CBO Budget Distribution",  "⚠️ Partial",       "CBO doesn't expose planned allocation",              "Engineering", "Low"],
    ["Ad Sets Paused / Killed",  "❌ Unavailable",   "No adset status history in database",                "Engineering", "Medium"],
]

tbl = doc.add_table(rows=1 + len(rows), cols=5)
tbl.style     = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

col_widths = [Inches(1.6), Inches(1.1), Inches(2.2), Inches(1.1), Inches(0.8)]

# Header row
for ci, hdr in enumerate(headers):
    cell = tbl.rows[0].cells[ci]
    cell.width = col_widths[ci]
    set_cell_bg(cell, "1877F2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(hdr)
    r.font.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
row_colors = ["FFFFFF", "F9FAFB"]
for ri, row_data in enumerate(rows, start=1):
    bg = row_colors[ri % 2]
    for ci, val in enumerate(row_data):
        cell = tbl.rows[ri].cells[ci]
        cell.width = col_widths[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(val)
        r.font.size = Pt(9)
        if ci == 1 and "Unavailable" in val:
            r.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
        elif ci == 1:
            r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        else:
            r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        if ci == 0:
            r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  NEXT STEPS
# ══════════════════════════════════════════════════════════════
add_para(doc, "Next Steps", size=14, bold=True, color="111827", space_before=8, space_after=8)

steps = [
    ("Media Buyer",  "Apply naming conventions (RET, BROAD, LAL, COLD) to all active ad sets — unlocks Retargeting CPP and Broad CPP immediately with no engineering work."),
    ("Engineering",  "Implement adset status snapshotting in the sync pipeline — unlocks Ad Sets Paused / Killed."),
    ("Product",      "Decide whether CBO Distribution should be surfaced as 'Actual Spend Share' with a disclaimer, and schedule accordingly."),
]

for i, (owner, desc) in enumerate(steps, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{i}.  [{owner}]  ")
    r1.font.size = Pt(10); r1.font.bold = True; r1.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

doc.add_paragraph()
add_rule(doc)
add_para(doc,
    "This document will be updated as fields are resolved and moved to active status.",
    size=9, color="9CA3AF", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)

# ── Save ──────────────────────────────────────────────────────
out_path = "C:/Users/Karan/Desktop/Projects/Digitally next_projects/adauto/docs/AdAuto_Unpopulated_Fields.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
